"""DOCX / TXT / HTML 보고서 생성.

polaris.py 에서 분리된 보고서 모듈:
  - 데이터 수집  : _report_collect (K8sManager 메서드들을 묶어 한 번에 조회)
  - 평가/finding : _report_evaluate (이상 징후 + 권장사항 산출)
  - DOCX 작성   : _report_write_docx (가장 큼 — 표지/요약/상세/AI 분석)
  - TXT 작성    : _report_write_txt
  - HTML 보고서 : _build_report_html (요약 카드 + 표)
  - LLM 호출    : _llm_ask (외부 OpenAI 호환 엔드포인트)

PolarisAPI.start_report() 가 이 모듈의 함수들을 백그라운드 스레드에서 호출.
"""
import json
import urllib.request
import urllib.error
from datetime import datetime
from pathlib import Path

from src.k8s import K8sManager, _age, _parse_cpu, _parse_mem, _pct

# DOCX 의존 — frozen build 에서 hiddenimports 로 끌어옴
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ─────────────────────────────────────────────────────────────────────────────

try:
    from docx import Document as _DocxDoc
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OxmlElem
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

def _report_required_dataset_keys() -> list[str]:
    return [
        'nodes', 'node_metrics', 'namespaces', 'pods', 'pod_metrics',
        'deployments', 'statefulsets', 'daemonsets', 'jobs', 'cronjobs',
        'services', 'ingresses', 'ingress_classes',
        'pvcs', 'pvs', 'configmaps', 'secrets', 'helm',
        'hpa', 'pdbs', 'network_policies', 'storage_classes',
        'resource_quotas', 'limit_ranges', 'rbac', 'kube_system', 'events',
    ]


_REPORT_SEVERITY_LEVEL = {
    'critical': 'CRITICAL',
    'high': 'HIGH',
    'medium': 'WARNING',
    'low': 'INFO',
    'info': 'INFO',
}


def _report_org():
    """보고서 작성 조직명. 이 빌드에서는 조직 정보를 표기하지 않는다 (None)."""
    return None


def _report_finding(category: str, severity: str, namespace: str, name: str,
                    value: str, detail: str, rec: str) -> dict:
    severity = (severity or 'info').lower()
    level = _REPORT_SEVERITY_LEVEL.get(severity, 'INFO')
    namespace = namespace or 'cluster'
    name = name or '-'
    value = str(value or '-')
    detail = str(detail or '')
    rec = str(rec or '')
    message = f'{namespace}/{name}: {value} — {detail}'
    if rec:
        message += f' | 권장: {rec}'
    return {
        'level': level,
        'severity': severity,
        'category': category,
        'namespace': namespace,
        'name': name,
        'value': value,
        'detail': detail,
        'rec': rec,
        'message': message,
    }


def _report_build_priority_summary(findings: list[dict]) -> dict:
    counts = {'critical': 0, 'high': 0, 'medium': 0, 'low': 0, 'info': 0}
    summary = {'counts': counts, 'immediate': [], 'short_term': [], 'backlog': []}
    for finding in findings or []:
        sev = (finding.get('severity') or '').lower()
        if not sev:
            level = (finding.get('level') or 'INFO').upper()
            sev = {'CRITICAL': 'critical', 'HIGH': 'high',
                   'WARNING': 'medium', 'INFO': 'info'}.get(level, 'info')
        counts.setdefault(sev, 0)
        counts[sev] += 1
        item = {
            'severity': sev,
            'category': finding.get('category', ''),
            'target': f"{finding.get('namespace', 'cluster')}/{finding.get('name', '-')}",
            'value': finding.get('value', ''),
            'detail': finding.get('detail') or finding.get('message', ''),
            'recommendation': finding.get('rec', ''),
        }
        if sev in ('critical', 'high'):
            summary['immediate'].append(item)
        elif sev == 'medium':
            summary['short_term'].append(item)
        else:
            summary['backlog'].append(item)
    return summary


def _report_collect(k8s, log_fn):
    """K8s 전체 데이터 수집."""
    data: dict = {}
    steps = [
        ('nodes',            k8s.get_node_extended),
        ('node_metrics',     k8s.get_node_metrics),
        ('namespaces',       k8s.get_namespaces_extended),
        ('pods',             lambda: k8s.get_pods(None)),
        ('pod_metrics',      k8s.get_pod_metrics_all),
        ('deployments',      k8s.get_deployments_extended),
        ('statefulsets',     k8s.get_statefulsets_extended),
        ('daemonsets',       k8s.get_daemonsets_extended),
        ('jobs',             lambda: k8s.get_jobs(None)),
        ('cronjobs',         lambda: k8s.get_cronjobs(None)),
        ('services',         lambda: k8s.get_services(None)),
        ('ingresses',        lambda: k8s.get_ingresses(None)),
        ('ingress_classes',  k8s.get_ingress_classes),
        ('pvcs',             lambda: k8s.get_pvcs(None)),
        ('pvs',              k8s.get_pvs),
        ('configmaps',       lambda: k8s.get_configmaps(None)),
        ('secrets',          lambda: k8s.get_secrets(None)),
        ('helm',             k8s.get_helm_releases),
        ('hpa',              k8s.get_hpa_extended),
        ('pdbs',             k8s.get_pdbs),
        ('network_policies', k8s.get_network_policies),
        ('storage_classes',  k8s.get_storage_classes),
        ('resource_quotas',  k8s.get_resource_quotas),
        ('limit_ranges',     k8s.get_limit_ranges),
        ('rbac',             k8s.get_rbac_summary),
        ('rbac_risky',       k8s.get_rbac_risky_subjects),   # 과도 권한 SA (v1.2.2)
        ('kube_system',      k8s.get_kube_system_info),
    ]
    for key, fn in steps:
        log_fn(f'  수집: {key}')
        try:
            data[key] = fn()
        except Exception as e:
            data[key] = [] if key not in ('rbac', 'kube_system') else {}
            log_fn(f'  [경고] {key} 수집 실패: {e}')
    data['hpas'] = data.get('hpa') or []
    log_fn('  수집: events')
    try:
        events = k8s.core.list_event_for_all_namespaces(
            field_selector='type=Warning', limit=150).items
        data['events'] = [{
            'namespace': e.metadata.namespace or '',
            'type': e.type or 'Warning',
            'reason': e.reason or '',
            'obj': f'{e.involved_object.kind}/{e.involved_object.name}',
            'message': (e.message or '')[:220],
            'count': e.count or 1,
            'last_time': str(e.last_timestamp)[:19] if e.last_timestamp else '-',
        } for e in events]
    except Exception as e:
        data['events'] = []
        log_fn(f'  [경고] events 수집 실패: {e}')
    data['cluster_version'] = k8s.cluster_info.get('version', 'N/A')
    data['collected_at'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    return data


def _report_evaluate(data):
    """규칙 기반 발견 사항 평가."""
    findings = []

    # 1. 재시작 과다 파드 / 비정상 파드
    for p in (data.get('pods') or []):
        try:
            r = int(p.get('restarts', 0))
        except (ValueError, TypeError):
            r = 0
        if r > 1000:
            sev, detail = 'critical', '즉시 분석 필요 (OOMKill / Liveness 실패 가능성)'
        elif r > 200:
            sev, detail = 'high', '반복 재시작 — 메모리 누수 또는 의존성 장애 가능성'
        elif r > 50:
            sev, detail = 'medium', '비정상 재시작 — 로그와 이벤트 모니터링 필요'
        elif r > 10:
            sev, detail = 'low', '간헐적 재시작'
        else:
            sev = detail = ''
        if sev:
            findings.append(_report_finding(
                'pod_restart', sev, p.get('namespace', ''), p.get('name', ''),
                f'{r}회', detail, 'kubectl logs / describe 로 원인 분석',
            ))

        status = p.get('status', '')
        if status and status not in ('Running', 'Succeeded', 'Completed'):
            findings.append(_report_finding(
                'pod_state', 'high', p.get('namespace', ''), p.get('name', ''),
                status, f'파드 비정상 상태: {status}', 'kubectl describe pod 로 이벤트 확인',
            ))

    # 2. NotReady 노드
    for n in (data.get('nodes') or []):
        for cond in (n.get('conditions') or []):
            if cond['type'] == 'Ready' and cond['status'] != 'True':
                findings.append(_report_finding(
                    'node_notready', 'critical', 'cluster', n.get('name', ''),
                    'NotReady', cond.get('reason') or 'Node Ready 조건 False',
                    'kubectl describe node 로 원인 확인',
                ))

    # 3. 노드 압박
    for n in (data.get('nodes') or []):
        for cond in (n.get('conditions') or []):
            if cond['type'] in ('MemoryPressure', 'DiskPressure', 'PIDPressure') and cond['status'] == 'True':
                findings.append(_report_finding(
                    'node_pressure', 'high', 'cluster', n.get('name', ''),
                    cond['type'], f"{cond['type']} 감지 — 리소스 임박",
                    '노드 리소스 확인 및 파드 재분산',
                ))

    # 4. NodePort 서비스
    np_svcs = [s for s in (data.get('services') or []) if s.get('type') == 'NodePort']
    sensitive = {'32000', '32016', '31043', '32033', '32198', '31000', '31634', '31633'}
    for svc in np_svcs:
        ports = str(svc.get('ports', ''))
        sev = 'high' if any(port in ports for port in sensitive) else 'medium'
        findings.append(_report_finding(
            'nodeport', sev, svc.get('namespace', ''), svc.get('name', ''),
            ports or 'NodePort', 'NodePort 직접 외부 노출',
            'Ingress + TLS/OAuth 또는 방화벽 ACL 적용 검토',
        ))

    # 5. HPA 없는 Deployment (replicas >= 2)
    hpa_items = data.get('hpa') or data.get('hpas') or []
    hpa_targets = {h.get('reference', '').split('/', 1)[-1] for h in hpa_items}
    no_hpa = [d for d in (data.get('deployments') or [])
              if d.get('desired', 0) >= 2 and d.get('name') not in hpa_targets]
    if len(no_hpa) > 3 or len(hpa_items) < 3:
        findings.append(_report_finding(
            'hpa', 'medium', 'cluster', 'HPA 설정',
            f'{len(hpa_items)}개', f'HPA 적용 워크로드 부족 — replicas≥2 미설정 {len(no_hpa)}개',
            '주요 서비스에 HPA 우선 적용',
        ))

    # 6. Lost / Pending PVC
    for pvc in (data.get('pvcs') or []):
        st = pvc.get('status', '')
        if st == 'Lost':
            findings.append(_report_finding(
                'pvc', 'critical', pvc.get('namespace', ''), pvc.get('name', ''),
                'Lost', 'PV 소실 — 데이터 잔류 가능성',
                '확인 후 안전 삭제 또는 백업 복구',
            ))
        elif st == 'Pending':
            findings.append(_report_finding(
                'pvc', 'medium', pvc.get('namespace', ''), pvc.get('name', ''),
                'Pending', 'PVC 바인딩 대기',
                'StorageClass, PV 용량, accessMode 확인',
            ))

    # 7. PDB disruptions_allowed == 0
    for pdb in (data.get('pdbs') or []):
        curr = pdb.get('current_healthy', 0)
        desired = pdb.get('desired_healthy', 0)
        unhealthy = isinstance(curr, int) and isinstance(desired, int) and curr < desired
        if pdb.get('disruptions_allowed', 1) == 0 or unhealthy:
            findings.append(_report_finding(
                'pdb', 'high' if unhealthy else 'medium',
                pdb.get('namespace', ''), pdb.get('name', ''),
                f'healthy {curr}/{desired}, disruptions {pdb.get("disruptions_allowed", 0)}',
                'PDB 조건 미충족 또는 중단 허용 0',
                '워크로드 복구 후 PDB 조건 재확인',
            ))

    # 8. StorageClass 기본값 없음
    scs = data.get('storage_classes') or []
    if scs and not any(s.get('is_default') == 'true' for s in scs):
        findings.append(_report_finding(
            'storageclass', 'medium', 'cluster', 'StorageClass',
            '기본값 없음', 'Default StorageClass 미설정 — 동적 PVC 프로비저닝 불가',
            '기본 StorageClass 지정 또는 PVC별 storageClassName 명시',
        ))

    # 9. kube-system 이상
    ks = data.get('kube_system') or {}
    for p in (ks.get('pods') or []):
        if p.get('status') not in ('Running', 'Succeeded', 'Completed') or (p.get('restarts') or 0) > 5:
            findings.append(_report_finding(
                'kube_system', 'critical' if p.get('status') not in ('Running', 'Succeeded', 'Completed') else 'medium',
                'kube-system', p.get('name', ''),
                f"{p.get('status', 'Unknown')}, restart {p.get('restarts', 0)}",
                'kube-system 파드 이상 — 클러스터 기능 저하 가능',
                'kubectl describe pod -n kube-system 확인',
            ))

    # 10. Deployment Ready < Desired
    for d in (data.get('deployments') or []):
        desired = d.get('desired', 0)
        ready = d.get('ready', 0)
        if desired > 0 and ready < desired:
            findings.append(_report_finding(
                'deployment', 'high', d.get('namespace', ''), d.get('name', ''),
                f'ready {ready}/{desired}', 'Deployment 레플리카 부족',
                'kubectl rollout status deployment 로 확인',
            ))

    # 11. RBAC 과도 권한 (cluster-admin / 와일드카드) — 보안 (v1.2.2)
    for r in (data.get('rbac_risky') or []):
        skind = r.get('subject_kind', '') or 'Subject'
        sname = r.get('name', '') or '-'
        subj  = f'{skind}/{sname}'
        # ServiceAccount 가 클러스터 관리자급 권한을 가지면 가장 위험(critical)
        sev = 'critical' if skind == 'ServiceAccount' else 'high'
        findings.append(_report_finding(
            'rbac', sev, r.get('namespace', '') or 'cluster', subj,
            r.get('reason', '과도한 권한'),
            f"{r.get('binding_kind', '')} '{r.get('binding', '')}' → "
            f"{r.get('role_kind', '')} '{r.get('role', '')}' (클러스터 전체 제어 가능)",
            '최소 권한 원칙에 따라 전용 Role 로 필요한 권한만 부여하세요',
        ))

    order = {'critical': 0, 'high': 1, 'medium': 2, 'low': 3, 'info': 4}
    findings.sort(key=lambda f: order.get(f.get('severity', 'info'), 9))

    return findings


def _llm_ask(url, model, user_prompt):
    """OpenAI 호환 API 단순 호출. max_tokens 생략 — 모델 기본값 사용."""
    import urllib.request
    from urllib.parse import urlparse as _urlparse

    # ── URL 스킴 검증 (보안): file://, ftp:// 등 비HTTP 스킴 차단 ──────────────
    _parsed = _urlparse(url)
    if _parsed.scheme not in ('http', 'https') or not _parsed.netloc:
        raise ValueError(
            f'허용되지 않는 LLM URL입니다: {url!r}\n'
            'http:// 또는 https:// 로 시작하는 URL만 사용할 수 있습니다.'
        )

    payload = json.dumps({
        'model': model,
        'messages': [
            {'role': 'system',
             'content': '당신은 Kubernetes 운영 전문가입니다. 한국어로 간결하게 답변하세요.'},
            {'role': 'user', 'content': user_prompt},
        ],
        'temperature': 0.3,
    }).encode('utf-8')
    req = urllib.request.Request(
        url.rstrip('/') + '/v1/chat/completions',
        data=payload,
        headers={'Content-Type': 'application/json'},
        method='POST',
    )
    with urllib.request.urlopen(req, timeout=120) as resp:
        result = json.loads(resp.read())
    return result['choices'][0]['message']['content'].strip()


def _set_cell_bg(cell, hex_color):
    """python-docx 셀 배경색 설정 (XML 직접 조작)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = _OxmlElem('w:shd')
    shd.set(_qn('w:val'), 'clear')
    shd.set(_qn('w:color'), 'auto')
    shd.set(_qn('w:fill'), hex_color)
    tcPr.append(shd)


def _report_write_docx(data, findings, llm_fn, out_path, log_fn):
    """v2 스타일 DOCX 보고서 — v1 자동 생성 가능 항목 포함."""
    if not HAS_DOCX:
        log_fn('  [경고] python-docx 미설치 → TXT 폴백')
        txt_path = str(out_path).replace('.docx', '.txt')
        _report_write_txt(data, findings, '', txt_path)
        return txt_path

    # ── 팔레트 ────────────────────────────────────────────────────────────────
    NAVY   = RGBColor(0x1F, 0x38, 0x64)
    BLUE   = RGBColor(0x2E, 0x75, 0xB6)
    RED    = RGBColor(0xC0, 0x00, 0x00)
    ORANGE = RGBColor(0xFF, 0x7F, 0x27)
    GREEN  = RGBColor(0x38, 0x86, 0x38)
    WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
    DARK   = RGBColor(0x17, 0x17, 0x17)
    GRAY   = RGBColor(0x76, 0x76, 0x76)
    AMBER  = RGBColor(0xB8, 0x86, 0x00)

    HEX_NAVY  = '1F3864'
    HEX_CRITD = 'F4CCCC'
    HEX_HIGHD = 'FCE4D6'
    HEX_MEDD  = 'FFF2CC'
    HEX_LOWD  = 'E2EFDA'
    HEX_LGRAY = 'F2F2F2'
    FONT = '맑은 고딕'

    SEV_KO  = {'critical': '긴급', 'high': '높음', 'medium': '중간',
               'warning': '중간', 'low': '낮음', 'info': '정보'}
    SEV_ORD = {'긴급': 0, '높음': 1, '중간': 2, '낮음': 3, '정보': 4}

    def sev_ko(sev):
        return SEV_KO.get(str(sev or '').lower(), '정보')

    def sev_bg(sev):
        k = sev_ko(sev)
        return (HEX_CRITD if k == '긴급' else HEX_HIGHD if k == '높음'
                else HEX_MEDD if k == '중간' else HEX_LOWD)

    def sev_rgb(sev):
        k = sev_ko(sev)
        return (RED if k == '긴급' else ORANGE if k == '높음'
                else AMBER if k == '중간' else GREEN)

    # ── XML/폰트 헬퍼 ────────────────────────────────────────────────────────
    def _para_bg(para, hex_color):
        pPr = para._p.get_or_add_pPr()
        shd = _OxmlElem('w:shd')
        shd.set(_qn('w:val'), 'clear')
        shd.set(_qn('w:color'), 'auto')
        shd.set(_qn('w:fill'), hex_color)
        pPr.append(shd)

    def _rfont(run, size=10, bold=False, color=None):
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color

    def _fix_ea(doc):
        def fix(run):
            rPr = run._r.get_or_add_rPr()
            rf = rPr.find(_qn('w:rFonts'))
            if rf is None:
                rf = _OxmlElem('w:rFonts')
                rPr.insert(0, rf)
            rf.set(_qn('w:eastAsia'), FONT)
        for p in doc.paragraphs:
            for r in p.runs: fix(r)
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for r in p.runs: fix(r)

    # ── 문서 초기화 ───────────────────────────────────────────────────────────
    doc = _DocxDoc()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.0)

    def _setup_hf():
        for sec in doc.sections:
            sec.header_distance = Cm(1.0)
            hdr = sec.header
            hdr.is_linked_to_previous = False
            hp = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
            hp.clear(); hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = hp.add_run('Kubernetes 클러스터 운영 점검 리포트')
            _rfont(r, size=9, color=GRAY)
            pPr = hp._p.get_or_add_pPr()
            pBdr = _OxmlElem('w:pBdr')
            bot = _OxmlElem('w:bottom')
            bot.set(_qn('w:val'), 'single'); bot.set(_qn('w:sz'), '4')
            bot.set(_qn('w:space'), '1');    bot.set(_qn('w:color'), 'CCCCCC')
            pBdr.append(bot); pPr.append(pBdr)
            sec.footer_distance = Cm(1.0)
            ftr = sec.footer
            ftr.is_linked_to_previous = False
            fp = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
            fp.clear(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_pre = fp.add_run('- '); _rfont(r_pre, size=9, color=GRAY)
            fp_p = fp._p
            rb = _OxmlElem('w:r'); fb = _OxmlElem('w:fldChar')
            fb.set(_qn('w:fldCharType'), 'begin'); rb.append(fb); fp_p.append(rb)
            ri = _OxmlElem('w:r'); it = _OxmlElem('w:instrText')
            it.text = ' PAGE '; ri.append(it); fp_p.append(ri)
            re_ = _OxmlElem('w:r'); fe = _OxmlElem('w:fldChar')
            fe.set(_qn('w:fldCharType'), 'end'); re_.append(fe); fp_p.append(re_)
            r_post = fp.add_run(' -'); _rfont(r_post, size=9, color=GRAY)

    _setup_hf()

    # ── DocBuilder ────────────────────────────────────────────────────────────
    def h1(text):
        p = doc.add_paragraph(); p.clear()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
        _para_bg(p, HEX_NAVY)
        pPr = p._p.get_or_add_pPr()
        ind = _OxmlElem('w:ind'); ind.set(_qn('w:left'), '100'); pPr.append(ind)
        run = p.add_run(f'  {text}')
        _rfont(run, size=13, bold=True, color=WHITE)

    def h2(text):
        p = doc.add_paragraph(); p.clear()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(3)
        pPr = p._p.get_or_add_pPr()
        pBdr = _OxmlElem('w:pBdr')
        L = _OxmlElem('w:left')
        L.set(_qn('w:val'), 'single'); L.set(_qn('w:sz'), '18')
        L.set(_qn('w:space'), '4');    L.set(_qn('w:color'), '2E75B6')
        pBdr.append(L); pPr.append(pBdr)
        ind = _OxmlElem('w:ind'); ind.set(_qn('w:left'), '140'); pPr.append(ind)
        run = p.add_run(text)
        _rfont(run, size=11, bold=True, color=NAVY)

    def para(text, bold=False, color=None, size=10, indent=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        if indent: p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(text)
        _rfont(run, size=size, bold=bold, color=color)

    def bullet(text, color=None, size=9.5):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Cm(0.5)
        run = p.add_run(f'• {text}')
        _rfont(run, size=size, color=color)

    def callout(text, hex_bg, color=None):
        p = doc.add_paragraph(); p.clear()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        pPr = p._p.get_or_add_pPr()
        ind = _OxmlElem('w:ind')
        ind.set(_qn('w:left'), '120'); ind.set(_qn('w:right'), '120')
        pPr.append(ind)
        _para_bg(p, hex_bg)
        run = p.add_run(text)
        _rfont(run, size=9.5, color=color or DARK)

    def tbl(headers, rows, col_widths=None, alt=True):
        if not rows:
            p = doc.add_paragraph('(데이터 없음)')
            if p.runs: p.runs[0].font.italic = True; p.runs[0].font.size = Pt(9)
            return
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = 'Table Grid'
        hdr = t.rows[0]
        for i, h in enumerate(headers):
            c = hdr.cells[i]; c.text = ''
            r = c.paragraphs[0].add_run(h)
            _rfont(r, size=9, bold=True, color=WHITE)
            _set_cell_bg(c, HEX_NAVY)
            if col_widths and i < len(col_widths): c.width = Cm(col_widths[i])
        for ridx, row_data in enumerate(rows):
            tr = t.add_row()
            row_bg = HEX_LGRAY if (alt and ridx % 2 == 0) else 'FFFFFF'
            for ci, val in enumerate(row_data):
                c = tr.cells[ci]; c.text = ''
                if isinstance(val, tuple):
                    txt     = str(val[0]) if val[0] is not None else ''
                    vbold   = val[1] if len(val) > 1 else False
                    vclr    = val[2] if len(val) > 2 else None
                    valg    = val[3] if len(val) > 3 else 'LEFT'
                    cell_bg = val[4] if len(val) > 4 else row_bg
                else:
                    txt, vbold, vclr, valg, cell_bg = str(val or ''), False, None, 'LEFT', row_bg
                r = c.paragraphs[0].add_run(txt)
                _rfont(r, size=9, bold=vbold, color=vclr)
                _set_cell_bg(c, cell_bg)
                pa = c.paragraphs[0]
                if valg in ('CENTER', 'center'): pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif valg in ('RIGHT', 'right'): pa.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                if col_widths and ci < len(col_widths): c.width = Cm(col_widths[ci])
        doc.add_paragraph('')

    def spacer(n=1):
        for _ in range(n):
            p = doc.add_paragraph('')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)

    def ai_comment(label, ctx):
        """섹션 AI 한 마디 — llm_fn 없으면 no-op."""
        if not llm_fn:
            return
        try:
            log_fn(f'    [AI] {label}...')
            text = llm_fn(ctx)
            if not text or not text.strip():
                return
            spacer()
            p = doc.add_paragraph()
            run = p.add_run('▶ AI: ' + text.strip())
            _rfont(run, size=9.5, color=BLUE)
            run.font.italic = True
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(8)
        except Exception as _e:
            log_fn(f'    [AI] {label} 실패: {_e}')

    # ── 기초 데이터 & 분석 ────────────────────────────────────────────────────
    priority  = _report_build_priority_summary(findings)
    counts    = priority['counts']
    nodes_all = data.get('nodes') or []
    pods_all  = data.get('pods')  or []
    pods_nr   = [p for p in pods_all
                 if p.get('status') not in ('Running', 'Succeeded', 'Completed')]
    pods_hi   = sorted([p for p in pods_all if int(p.get('restarts', 0)) > 5],
                       key=lambda x: int(x.get('restarts', 0)), reverse=True)[:20]
    svc_all   = data.get('services') or []
    nodeport  = [s for s in svc_all if s.get('type') == 'NodePort']
    pvcs_all  = data.get('pvcs') or []
    pvcs_lost = [p for p in pvcs_all if p.get('status') not in ('Bound', 'Released')]
    deps_all  = data.get('deployments') or []
    sts_all   = data.get('statefulsets') or []
    hpas      = data.get('hpa') or []
    ingresses = data.get('ingresses') or []
    nss       = data.get('namespaces') or []
    helmlist  = data.get('helm') or []

    ko_counts = {'긴급': counts.get('critical', 0),
                 '높음': counts.get('high', 0),
                 '중간': counts.get('medium', 0) + counts.get('warning', 0),
                 '낮음': counts.get('low', 0),
                 '정보': counts.get('info', 0)}

    # 노드 역할 분석
    def _parse_roles(n):
        return [r.strip() for r in n.get('roles', 'worker').split(',') if r.strip()]

    all_roles = set()
    for n in nodes_all:
        for r in _parse_roles(n): all_roles.add(r)
    has_dedicated_worker = any(
        not any(cp in _parse_roles(n) for cp in ('control-plane', 'master', 'etcd'))
        for n in nodes_all)
    kernels = list({n.get('kernel', '') for n in nodes_all if n.get('kernel', '')})
    kernel_diff = len(kernels) > 1

    # Deployment/StatefulSet 이상
    deps_nr = [d for d in deps_all
               if int(d.get('ready', 0)) < int(d.get('desired', 0))]
    sts_nr  = [s for s in sts_all
               if int(s.get('ready', 0)) < int(s.get('desired', 0))]

    # HPA 미적용 워크로드
    hpa_targets = set()
    for h in hpas:
        ref = h.get('reference', '')
        parts = ref.split('/')
        if len(parts) >= 2: hpa_targets.add(parts[-1].lower())
    deps_no_hpa = [d for d in deps_all if d.get('name', '').lower() not in hpa_targets]

    # Ingress class 분류
    from collections import Counter as _Counter
    class_dist = _Counter(i.get('class', '<none>') or '<none>' for i in ingresses)
    no_class_ns = sorted({i.get('namespace', '')
                          for i in ingresses
                          if (i.get('class', '<none>') or '<none>') == '<none>'})

    # 대용량 PVC (50Gi 이상)
    def _parse_gi(cap):
        s = str(cap or '').upper().strip()
        try:
            if 'TI' in s: return float(s.replace('TI', '').strip()) * 1024
            if 'GI' in s: return float(s.replace('GI', '').strip())
            if 'G' in s:  return float(s.replace('G', '').strip())
        except: pass
        return 0
    large_pvcs = sorted(
        [(p, _parse_gi(p.get('capacity', ''))) for p in pvcs_all
         if _parse_gi(p.get('capacity', '')) >= 50],
        key=lambda x: x[1], reverse=True)

    # ArgoCD 감지
    has_argocd = any(n.get('name', '') == 'argocd' for n in nss)
    # Helm 이상
    helm_failed = [h for h in helmlist if h.get('status', '') not in ('deployed', 'superseded')]

    # ═══════════════════════════════════════════════════════════════════════════
    # 표지
    # ═══════════════════════════════════════════════════════════════════════════
    log_fn('  [1/10] 표지...')
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(60)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Kubernetes 클러스터')
    _rfont(r, size=26, bold=True, color=NAVY)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('운영 점검 리포트')
    _rfont(r2, size=22, bold=True, color=NAVY)
    spacer()
    _org = _report_org()
    for txt, sz, clr in [row for row in [
        (f"점검 일시: {data.get('collected_at', '')}", 12, GRAY),
        (f'작성 조직: {_org}', 12, GRAY) if _org else None,
        (f"클러스터: {data.get('cluster_version', 'N/A')}", 11, GRAY),
        (f"노드: {len(nodes_all)}개", 11, GRAY),
    ] if row]:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(txt)
        _rfont(r3, size=sz, color=clr)
    spacer(2)
    chip_data = [('긴급', ko_counts['긴급'], HEX_CRITD, RED),
                 ('높음', ko_counts['높음'], HEX_HIGHD, ORANGE),
                 ('중간', ko_counts['중간'], HEX_MEDD,  AMBER),
                 ('낮음', ko_counts['낮음'], HEX_LOWD,  GREEN)]
    t_cov = doc.add_table(rows=2, cols=4)
    t_cov.style = 'Table Grid'
    t_cov.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ci, (label, cnt, bg, clr) in enumerate(chip_data):
        ch = t_cov.rows[0].cells[ci]; ch.text = ''
        rh = ch.paragraphs[0].add_run(label)
        _rfont(rh, size=10, bold=True, color=WHITE)
        ch.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(ch, HEX_NAVY); ch.width = Cm(4.125)
        cv = t_cov.rows[1].cells[ci]; cv.text = ''
        rv = cv.paragraphs[0].add_run(str(cnt))
        _rfont(rv, size=22, bold=True, color=clr)
        cv.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(cv, bg); cv.width = Cm(4.125)
    doc.add_paragraph('')
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 1. 점검 개요
    # ═══════════════════════════════════════════════════════════════════════════
    log_fn('  [2/10] 점검 개요...')
    h1('1. 점검 개요')
    h2('1.1 점검 목적')
    para('본 리포트는 운영 중인 Kubernetes 클러스터의 전반적인 상태를 점검하고, '
         '잠재적 위험 요소를 식별하여 안정적인 서비스 운영을 위한 개선 방향을 제시합니다.', size=10)
    h2('1.2 점검 범위')
    for item in ['클러스터 기본 현황 (노드, 버전, 리소스)',
                 '워크로드 상태 (Pod, Deployment, StatefulSet 등)',
                 'HPA 및 오토스케일링 구성',
                 'Ingress 및 Service 노출 구조',
                 'PVC / 스토리지 상태',
                 '보안 위험 항목 (RBAC, NetworkPolicy)',
                 '발견 이슈 분석 및 개선 우선순위']:
        bullet(item)
    h2('1.3 점검 정보')
    tbl(['항목', '내용'], [row for row in [
        ['클러스터 버전',   data.get('cluster_version', 'N/A')],
        ['점검 일시',       data.get('collected_at', '')],
        (['작성 조직', _org] if _org else None),
        ['노드 수',         f"{len(nodes_all)}개"],
        ['전체 파드 수',    f"{len(pods_all)}개"],
        ['네임스페이스 수', f"{len(nss)}개"],
        ['발견 이슈',       f"긴급 {ko_counts['긴급']}건 / 높음 {ko_counts['높음']}건 / "
                            f"중간 {ko_counts['중간']}건"],
    ] if row], col_widths=[4, 12.5])
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 2. 클러스터 기본 현황
    # ═══════════════════════════════════════════════════════════════════════════
    log_fn('  [3/10] 클러스터 기본 현황...')
    h1('2. 클러스터 기본 현황')

    h2('2.1 노드 상태')
    not_ready = [n for n in nodes_all if n.get('status') != 'Ready']
    ready_cnt = len(nodes_all) - len(not_ready)
    # 역할 분석 산문 (v1 스타일)
    role_desc = ', '.join(sorted(all_roles)) or 'worker'
    worker_note = (
        '' if has_dedicated_worker
        else f' 전용 워커 노드가 없어 시스템/사용자 워크로드가 동일 노드에서 실행됩니다.')
    para(f"총 {len(nodes_all)}개 노드 중 {ready_cnt}개가 Ready 상태이며, "
         f"{role_desc} 역할을 겸합니다.{worker_note}", size=10,
         color=RED if not_ready else None)
    if not_ready:
        callout(f"⚠ NotReady 노드: {', '.join(n.get('name','') for n in not_ready)}",
                HEX_CRITD, RED)
    if not has_dedicated_worker and len(nodes_all) > 0:
        callout('전용 워커 노드 없음 — 컨트롤 플레인 노드에 사용자 워크로드가 혼재합니다. '
                '장기적으로 전용 워커 노드 추가를 권장합니다.', HEX_MEDD, AMBER)
    if kernel_diff:
        callout(f"⚠ 커널 버전 혼재: {' / '.join(kernels)} — "
                f"노드 간 OS 패치 동기화가 필요합니다.", HEX_HIGHD, ORANGE)
    node_rows = []
    for n in nodes_all:
        st   = n.get('status', 'Unknown')
        ok   = st == 'Ready'
        conds = n.get('conditions') or []
        press = [c['type'].replace('Pressure', '') for c in conds
                 if c['type'] in ('MemoryPressure', 'DiskPressure', 'PIDPressure')
                 and c['status'] == 'True']
        node_rows.append([
            n.get('name', ''),
            n.get('roles', '-'),
            (st, True, GREEN if ok else RED, 'CENTER', HEX_LOWD if ok else HEX_CRITD),
            n.get('version', '-'),
            n.get('os', '-'),
            n.get('allocatable_cpu', '-'),
            n.get('allocatable_mem', '-'),
            n.get('kernel', '-'),
            ', '.join(press) or '정상',
        ])
    tbl(['노드명', '역할', '상태', '버전', 'OS', 'CPU(Alloc)', 'MEM(Alloc)', '커널', '압박'],
        node_rows, col_widths=[2.2, 1.5, 1.5, 2.0, 2.8, 1.7, 1.8, 1.5, 1.5])

    h2('2.2 전체 워크로드 현황')
    running   = len([p for p in pods_all if p.get('status') == 'Running'])
    completed = len([p for p in pods_all if p.get('status') in ('Succeeded', 'Completed')])
    tbl(['항목', '수치'], [
        ['전체 Pod',         str(len(pods_all))],
        ['Running',          str(running)],
        ['Completed (정상)', str(completed)],
        ['비정상',           str(len(pods_nr))],
        ['Deployment',       f"{len(deps_all)}개 (이상 {len(deps_nr)}개)"],
        ['StatefulSet',      f"{len(sts_all)}개 (이상 {len(sts_nr)}개)"],
        ['DaemonSet',        str(len(data.get('daemonsets') or []))],
        ['Job',              str(len(data.get('jobs') or []))],
        ['CronJob',          str(len(data.get('cronjobs') or []))],
        ['Service',          str(len(svc_all))],
        ['PVC',              str(len(pvcs_all))],
        ['Helm 릴리스',      f"{len(helmlist)}개 (이상 {len(helm_failed)}개)"],
    ], col_widths=[5, 11.5])

    h2('2.3 주요 관찰사항')
    obs = []
    if not deps_nr:
        obs.append(f"전체 Deployment {len(deps_all)}개 모두 READY 상태 (정상)")
    else:
        obs.append(f"Deployment {len(deps_nr)}개 미준비 — "
                   f"{', '.join(d.get('name','') for d in deps_nr[:5])}")
    if not sts_nr:
        obs.append(f"StatefulSet {len(sts_all)}개 모두 정상 운영 중")
    else:
        obs.append(f"StatefulSet {len(sts_nr)}개 미준비 — "
                   f"{', '.join(s.get('name','') for s in sts_nr[:3])}")
    if has_argocd:
        obs.append("ArgoCD GitOps 동기화 운영 중 (argocd 네임스페이스 확인)")
    if helm_failed:
        obs.append(f"Helm 비정상 릴리스 {len(helm_failed)}개 — "
                   f"{', '.join(h.get('name','') for h in helm_failed[:3])}")
    else:
        obs.append(f"Helm 릴리스 {len(helmlist)}개 모두 정상 배포 상태")
    ks = data.get('kube_system') or {}
    ks_pods = ks.get('pods') or []
    ks_bad  = [p for p in ks_pods
               if p.get('status') not in ('Running', 'Succeeded', 'Completed')]
    if ks_bad:
        obs.append(f"kube-system 비정상 파드 {len(ks_bad)}개 — 컨트롤 플레인 점검 필요")
    else:
        obs.append("kube-system 파드 정상 운영 중")
    for o in obs:
        bullet(o)

    h2('2.4 운영 중인 네임스페이스')
    ns_names = [n.get('name', '') for n in nss]
    para(f"총 {len(ns_names)}개 네임스페이스 운영 중:", size=10)
    # 한 행에 여러 이름 나열
    chunk = 8
    for i in range(0, len(ns_names), chunk):
        bullet(', '.join(ns_names[i:i+chunk]), size=9)
    ai_comment('클러스터 기본 현황', f"""Kubernetes 클러스터 기본 현황:
- 노드: 총 {len(nodes_all)}개 (Ready {len(nodes_all) - len(not_ready)}개, NotReady {len(not_ready)}개)
- 역할 구성: {', '.join(sorted(all_roles)) or '미확인'}
- 파드: {len(pods_all)}개 (비정상 {len(pods_nr)}개)
- Deployment {len(deps_all)}개, StatefulSet {len(sts_all)}개, 네임스페이스 {len(nss)}개

클러스터 전반 상태를 1~2문장으로 평가하세요.""")
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 3. 워크로드 상태
    # ═══════════════════════════════════════════════════════════════════════════
    log_fn('  [4/10] 워크로드 상태...')
    h1('3. 워크로드 상태')

    h2('3.1 비정상 파드')
    if pods_nr:
        callout(f"⚠ 비정상 파드 {len(pods_nr)}건 검출 — 즉각 원인 분석이 필요합니다.",
                HEX_CRITD, RED)
        tbl(['네임스페이스', '파드명', '위험도', '상태', '재시작', 'Age'],
            [[p.get('namespace', ''), p.get('name', ''),
              ('긴급', True, RED, 'CENTER', HEX_CRITD),
              p.get('status', ''), str(p.get('restarts', 0)), p.get('age', '-')]
             for p in pods_nr[:30]], col_widths=[2.5, 5, 1.5, 2.5, 1.5, 3.5])
    else:
        callout('✓ 비정상 파드 없음 — 전체 워크로드 Running/Completed 정상 상태입니다.',
                HEX_LOWD, GREEN)

    h2('3.2 재시작 이상 파드')
    if pods_hi:
        callout(f"⚠ 재시작 이상 파드 {len(pods_hi)}건 검출 — "
                "OOMKill, Liveness Probe 실패, 의존성 장애 등 원인 분석 필요",
                HEX_CRITD if any(int(p.get('restarts',0))>200 for p in pods_hi) else HEX_HIGHD,
                RED if any(int(p.get('restarts',0))>200 for p in pods_hi) else ORANGE)
        tbl(['네임스페이스', '파드명', '위험도', '재시작/운영', '상태'],
            [[p.get('namespace', ''), p.get('name', ''),
              ('긴급' if int(p.get('restarts',0)) > 200 else '높음',
               True,
               RED if int(p.get('restarts',0)) > 200 else ORANGE,
               'CENTER',
               HEX_CRITD if int(p.get('restarts',0)) > 200 else HEX_HIGHD),
              (f"{p.get('restarts',0)}회 ({p.get('age','-')})",
               True, RED if int(p.get('restarts',0)) > 200 else ORANGE),
              p.get('status', '-')]
             for p in pods_hi], col_widths=[2.5, 5.5, 1.5, 3.5, 3.5])
    else:
        callout('✓ 재시작 이상 파드 없음 — 전체 Pod 재시작 횟수가 정상 범위입니다.',
                HEX_LOWD, GREEN)

    if ks_bad:
        h2('3.3 kube-system 이상 파드')
        callout(f"⚠ kube-system 비정상 파드 {len(ks_bad)}개 — 컨트롤 플레인 점검 필요",
                HEX_CRITD, RED)
        tbl(['파드명', '상태', '재시작', '노드'],
            [[p.get('name',''), p.get('status',''),
              str(p.get('restarts',0)), p.get('node','-')]
             for p in ks_bad], col_widths=[6, 2.5, 2, 6])

    metrics = (data.get('pod_metrics') or [])[:10]
    if metrics:
        h2('3.4 메모리 사용 상위 파드 (Top 10)')
        para('다음 워크로드들은 메모리 사용량이 높아 Limit 미설정 시 OOMKill 또는 '
             '노드 메모리 부족 위험이 있습니다.', size=10)
        tbl(['네임스페이스', '파드명', 'CPU (m)', 'MEM (MiB)'],
            [[m.get('namespace',''), m.get('name',''),
              f"{m.get('cpu_m',0)}m", f"{m.get('mem_mi',0)}Mi"]
             for m in sorted(metrics, key=lambda x: int(x.get('mem_mi',0)), reverse=True)],
            col_widths=[3, 7, 2.5, 4])

    # Request/Limit 현황
    lr_items = data.get('limit_ranges') or []
    rq_items = data.get('resource_quotas') or []
    h2('3.5 리소스 Request/Limit 현황')
    lr_ns = {lr.get('namespace','') for lr in lr_items}
    rq_ns = {rq.get('namespace','') for rq in rq_items}
    para(f"LimitRange 설정 네임스페이스: {len(lr_ns)}개 | "
         f"ResourceQuota 설정 네임스페이스: {len(rq_ns)}개", size=10)
    if not lr_items:
        callout('LimitRange 미설정 — 네임스페이스에 LimitRange 적용 시 기본 Request/Limit 자동 부여 가능. '
                '미설정 Pod는 노드 전체 메모리를 소진할 수 있습니다.', HEX_MEDD, AMBER)
    else:
        bullet(f"LimitRange 적용 네임스페이스: {', '.join(sorted(lr_ns))}")
    ai_comment('워크로드 상태', f"""Kubernetes 워크로드 현황:
- 비정상 파드: {len(pods_nr)}개
- 재시작 5회 이상 파드: {len(pods_hi)}개
- LimitRange 적용 네임스페이스: {len(lr_ns)}개 / ResourceQuota: {len(rq_ns)}개

워크로드 상태의 원인 분석과 운영 조치를 1~2문장으로 작성하세요.""")
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 4. HPA / 오토스케일링
    # ═══════════════════════════════════════════════════════════════════════════
    log_fn('  [5/10] HPA/오토스케일링...')
    h1('4. HPA / 오토스케일링')

    h2('4.1 HPA 현황')
    para(f"점검 시점 기준 HPA(Horizontal Pod Autoscaler)가 설정된 워크로드는 "
         f"{len(hpas)}개입니다. 대부분의 워크로드가 고정 레플리카 수로 운영됩니다.", size=10)
    if not hpas:
        callout(f"현재 HPA {len(hpas)}개 — 트래픽 급증 시 수동 스케일 아웃만 가능합니다.",
                HEX_MEDD, AMBER)
    else:
        tbl(['네임스페이스', '이름', '대상', 'Min', 'Max', '현재'],
            [[h.get('namespace',''), h.get('name',''), h.get('reference','-'),
              str(h.get('min',1)), str(h.get('max',1)), str(h.get('current',0))]
             for h in hpas], col_widths=[2.5, 3.5, 4, 1.5, 1.5, 3.5])

    h2('4.2 HPA 미적용 주요 워크로드')
    # 표시 우선순위: 비정상 우선, 이후 이름순
    show_no_hpa = sorted(deps_no_hpa, key=lambda d: (
        0 if int(d.get('ready',0)) < int(d.get('desired',0)) else 1,
        d.get('name', '')))[:20]
    if show_no_hpa:
        para(f"HPA 미설정 Deployment {len(deps_no_hpa)}개 — "
             f"트래픽 급증 시 수동 스케일 아웃만 가능합니다.", size=10)
        tbl(['네임스페이스', '이름', 'Ready/Desired', '비고'],
            [[d.get('namespace',''), d.get('name',''),
              f"{d.get('ready',0)}/{d.get('desired',0)}",
              'HPA 도입 검토 권장']
             for d in show_no_hpa], col_widths=[3, 5.5, 2.5, 5.5])
    else:
        callout('✓ 모든 Deployment에 HPA가 적용되어 있습니다.', HEX_LOWD, GREEN)

    pdbs = data.get('pdbs') or []
    if pdbs:
        h2('4.3 PDB (Pod Disruption Budget)')
        tbl(['네임스페이스', '이름', 'MinAvail', 'CurrHealthy', 'DesiredHealthy'],
            [[p.get('namespace',''), p.get('name',''),
              str(p.get('min_available','-')), str(p.get('current_healthy',0)),
              str(p.get('desired_healthy',0))]
             for p in pdbs], col_widths=[3, 4, 2.5, 3, 4])
    ai_comment('HPA/오토스케일링', f"""HPA(오토스케일링) 현황:
- 설정된 HPA: {len(hpas)}개
- Deployment 중 HPA 미적용: {len(deps_no_hpa)}개 / 전체 {len(deps_all)}개
- StatefulSet: {len(sts_all)}개

HPA 적용 현황과 오토스케일링 운영 전략을 1~2문장으로 평가하세요.""")
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 5. Ingress / Service 구조
    # ═══════════════════════════════════════════════════════════════════════════
    log_fn('  [6/10] Ingress/Service...')
    h1('5. Ingress / Service 구조')

    h2('5.1 Ingress 현황')
    para(f"총 {len(ingresses)}개 Ingress 구성.", size=10)
    if class_dist:
        tbl(['IngressClass', '개수', '해당 네임스페이스 (예시)'],
            [[cls, str(cnt),
              ', '.join(sorted({i.get('namespace','') for i in ingresses
                                if (i.get('class','<none>') or '<none>') == cls})[:5])]
             for cls, cnt in class_dist.most_common()],
            col_widths=[3.5, 2, 11])
    if no_class_ns:
        callout(f"<none> IngressClass {class_dist.get('<none>',0)}개 — "
                f"명시적 class 지정 권고: "
                f"{', '.join(no_class_ns[:10])}"
                + ('...' if len(no_class_ns) > 10 else ''),
                HEX_MEDD, AMBER)

    h2('5.2 TLS/HTTPS 현황')
    # cert-manager 감지
    has_cert_manager = any(n.get('name','') == 'cert-manager' for n in nss)
    if has_cert_manager:
        para('cert-manager를 통한 TLS 인증서 자동 갱신이 운영 중입니다.', size=10)
        bullet('cert-manager 네임스페이스 확인됨 — Let\'s Encrypt 또는 내부 CA 연동 가능')
    tls_ingress  = [i for i in ingresses if i.get('hosts', '') not in ('*', '') and 'tls' in str(i)]
    no_tls       = [i for i in ingresses if i.get('address', '') and 'tls' not in str(i)]
    if no_tls:
        callout(f"HTTPS 미적용 Ingress {len(no_tls)}개 — TLS 전환 검토 필요", HEX_MEDD, AMBER)

    h2('5.3 NodePort 노출 서비스')
    if nodeport:
        callout(f"NodePort 직접 노출 서비스 {len(nodeport)}건 — "
                "방화벽 ACL 또는 Ingress 전환 필요", HEX_HIGHD, ORANGE)
        def _ports(s):
            p = str(s.get('ports') or '-')
            return p if len(p) <= 30 else p[:29] + '…'
        tbl(['네임스페이스', '서비스명', '위험도', '포트', '권장 조치'],
            [[s.get('namespace',''), s.get('name',''),
              ('높음', True, ORANGE, 'CENTER', HEX_HIGHD),
              _ports(s), '민감 서비스: Ingress + 방화벽 ACL 적용']
             for s in nodeport], col_widths=[2.5, 3.5, 1.5, 4, 5])
    else:
        callout('✓ NodePort 직접 노출 서비스 없음', HEX_LOWD, GREEN)

    lb_svcs = [s for s in svc_all if s.get('type') == 'LoadBalancer']
    if lb_svcs:
        h2('5.4 LoadBalancer 서비스')
        para(f"외부 노출 LoadBalancer 서비스 {len(lb_svcs)}개 — 보안 검토 권장.", size=10, color=ORANGE)
        def _ports2(s):
            p = str(s.get('ports') or '-')
            return p if len(p) <= 22 else p[:21] + '…'
        tbl(['네임스페이스', '서비스명', 'Cluster-IP', '포트', 'Age'],
            [[s.get('namespace',''), s.get('name',''),
              s.get('cluster-ip','-'), _ports2(s), s.get('age','-')]
             for s in lb_svcs], col_widths=[2.5, 4.5, 2.5, 4, 3])
    ai_comment('Ingress/Service 구조', f"""네트워크 노출 구조:
- Ingress: {len(ingresses)}개, Class 분포: {dict(class_dist)}
- NodePort 서비스: {len(nodeport)}개

네트워크 노출 관점의 보안·운영 이슈를 1~2문장으로 평가하세요.""")
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 6. PVC / 스토리지
    # ═══════════════════════════════════════════════════════════════════════════
    log_fn('  [7/10] PVC/스토리지...')
    h1('6. PVC / 스토리지')

    h2('6.1 스토리지 구성 개요')
    pvcs_bound = len([p for p in pvcs_all if p.get('status') == 'Bound'])
    sc_list    = data.get('storage_classes') or []
    sc_names   = [s.get('name','') for s in sc_list]
    para(f"주력 StorageClass: {', '.join(sc_names) if sc_names else 'N/A'}. "
         f"전체 PVC {len(pvcs_all)}개 중 Bound(정상) {pvcs_bound}개, "
         f"이상 {len(pvcs_lost)}개입니다.", size=10)
    nfs_only = all('nfs' in s.get('provisioner','').lower() for s in sc_list) and sc_list
    if nfs_only:
        callout('NFS 기반 StorageClass만 운영 중 — NFS 서버 가용성이 전체 스토리지의 '
                '단일 장애점(SPOF)입니다. 디스크 사용률 모니터링을 강화하세요.',
                HEX_MEDD, AMBER)
    tbl(['항목', '수치'], [
        ['전체 PVC',       str(len(pvcs_all))],
        ['Bound (정상)',   str(pvcs_bound)],
        ['Lost / Pending', str(len(pvcs_lost))],
        ['StorageClass 수', str(len(sc_list))],
    ], col_widths=[5, 11.5])

    h2('6.2 이상 PVC (Lost / Pending)')
    if pvcs_lost:
        callout(f"⚠ 이상 PVC {len(pvcs_lost)}건 — PV 소실 상태, 데이터 잔류 여부 즉시 확인 필요",
                HEX_CRITD, RED)
        tbl(['네임스페이스', 'PVC명', '위험도', '상태', '용량', '권장 조치'],
            [[p.get('namespace',''), p.get('name',''),
              ('긴급' if p.get('status')=='Lost' else '높음',
               True,
               RED if p.get('status')=='Lost' else ORANGE,
               'CENTER',
               HEX_CRITD if p.get('status')=='Lost' else HEX_HIGHD),
              p.get('status',''), p.get('capacity','-'),
              '데이터 복구 불필요 확인 후 안전 삭제']
             for p in pvcs_lost], col_widths=[2.5, 3.5, 1.5, 1.5, 2, 5.5])
    else:
        callout('✓ 이상 PVC 없음 — 전체 PVC가 정상 Bound 상태입니다.', HEX_LOWD, GREEN)

    if large_pvcs:
        h2(f'6.3 대용량 PVC ({len(large_pvcs)}개, 50Gi 이상)')
        tbl(['네임스페이스', 'PVC명', '용량', '상태', '비고'],
            [[p.get('namespace',''), p.get('name',''), p.get('capacity','-'),
              p.get('status','-'), 'retention 정책 점검 권장']
             for p, gi in large_pvcs[:15]], col_widths=[2.5, 5, 2, 2, 5])

    if sc_list:
        h2('6.4 StorageClass')
        tbl(['이름', 'Provisioner', 'Reclaim', '기본값'],
            [[s.get('name',''), s.get('provisioner',''),
              s.get('reclaim_policy','-'), s.get('is_default','false')]
             for s in sc_list], col_widths=[4, 7, 2.5, 3])
    ai_comment('PVC/스토리지', f"""스토리지 현황:
- PVC 총 {len(pvcs_all)}개 (Bound: {pvcs_bound}개, 이상: {len(pvcs_lost)}개)
- StorageClass: {', '.join(sc_names) if sc_names else 'N/A'}
- NFS 단독 구성: {'예' if nfs_only else '아니오'}

스토리지 운영 관점의 위험 요소를 1~2문장으로 평가하세요.""")
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 7. 보안
    # ═══════════════════════════════════════════════════════════════════════════
    log_fn('  [8/10] 보안...')
    h1('7. 보안')

    h2('7.1 RBAC 구성')
    rbac = data.get('rbac') or {}
    tbl(['항목', '수량'], [
        ['ClusterRole',        str(rbac.get('cluster_roles', 0))],
        ['ClusterRoleBinding', str(rbac.get('cluster_role_bindings', 0))],
        ['Role',               str(rbac.get('roles', 0))],
        ['RoleBinding',        str(rbac.get('role_bindings', 0))],
        ['ServiceAccount',     str(rbac.get('service_accounts', 0))],
    ], col_widths=[5, 11.5])

    h2('7.2 NetworkPolicy')
    netpols = data.get('network_policies') or []
    if not netpols:
        callout('ℹ NetworkPolicy 미설정 — 네임스페이스 간 트래픽 제한이 없습니다. '
                '주요 서비스 네임스페이스에 NetworkPolicy 적용을 권고합니다.', HEX_MEDD, DARK)
    else:
        para(f"총 {len(netpols)}개 NetworkPolicy 설정됨.", size=10)
    ai_comment('보안 점검', f"""보안 점검 결과:
- NodePort 노출 서비스: {len(nodeport)}개
- RBAC: ClusterRole {rbac.get('cluster_roles', 0)}개, ClusterRoleBinding {rbac.get('cluster_role_bindings', 0)}개
- NetworkPolicy: {len(netpols)}개

보안 관점의 핵심 위험 요소와 권고 조치를 1~2문장으로 작성하세요.""")
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 8. 발견 사항 및 권고
    # ═══════════════════════════════════════════════════════════════════════════
    log_fn('  [9/10] 발견 사항...')
    h1('8. 발견 사항 및 권고')

    if not findings:
        callout('✓ 발견된 이슈가 없습니다. 클러스터 상태가 양호합니다.', HEX_LOWD, GREEN)
    else:
        if ko_counts['긴급'] > 0:
            callout(f"⚠ 긴급 이슈 {ko_counts['긴급']}건: 즉각 조치 없으면 서비스 중단 또는 보안 사고 가능성이 높습니다.",
                    HEX_CRITD, RED)
        if ko_counts['높음'] > 0:
            callout(f"▲ 높음 이슈 {ko_counts['높음']}건: 단기 방치 시 장애로 발전할 수 있는 전조 증상입니다.",
                    HEX_HIGHD, ORANGE)
        sf = sorted(findings, key=lambda f: SEV_ORD.get(
            sev_ko(f.get('severity') or f.get('level', 'INFO')), 9))

        def sev_cell(f):
            orig = f.get('severity') or f.get('level', '')
            s = sev_ko(orig)
            return (s, True, sev_rgb(orig), 'CENTER', sev_bg(orig))

        tbl(['위험도', '유형', '대상', '현황', '설명', '권장 조치'],
            [[sev_cell(f), f.get('category',''),
              f"{f.get('namespace','cluster')}/{f.get('name','-')}",
              f.get('value',''), f.get('detail',''), f.get('rec','')]
             for f in sf], col_widths=[1.5, 2.5, 2.5, 2, 4, 4])

    h2('8.1 긴급 — 즉시 조치 (1주 이내)')
    if priority['immediate']:
        tbl(['#', '위험도', '대상', '현황', '권장 조치'],
            [[str(i+1),
              (sev_ko(r['severity']), True, sev_rgb(r['severity']), 'CENTER', sev_bg(r['severity'])),
              r['target'], r['value'], r['recommendation']]
             for i, r in enumerate(priority['immediate'][:15])],
            col_widths=[0.8, 1.5, 3.5, 2, 8.7])
    else:
        callout('✓ 즉시 조치가 필요한 긴급/높음 이슈가 없습니다.', HEX_LOWD, GREEN)

    h2('8.2 중간 — 1~4주 내 조치')
    if priority['short_term']:
        tbl(['#', '위험도', '대상', '현황', '권장 조치'],
            [[str(i+1),
              (sev_ko(r['severity']), True, sev_rgb(r['severity']), 'CENTER', sev_bg(r['severity'])),
              r['target'], r['value'], r['recommendation']]
             for i, r in enumerate(priority['short_term'][:15])],
            col_widths=[0.8, 1.5, 3.5, 2, 8.7])
    else:
        para('중간 우선순위 항목이 없습니다.', size=10, color=GRAY)

    if priority.get('backlog'):
        h2('8.3 낮음 — 분기 내 추적')
        tbl(['#', '위험도', '대상', '설명'],
            [[str(i+1),
              (sev_ko(r['severity']), True, sev_rgb(r['severity']), 'CENTER', sev_bg(r['severity'])),
              r['target'], r.get('detail', '-')]
             for i, r in enumerate(priority['backlog'][:12])],
            col_widths=[0.8, 1.5, 4, 10.2])

    events = data.get('events') or []
    if events:
        h2('8.4 Warning 이벤트')
        para(f"총 {len(events)}건의 Warning 이벤트가 수집되었습니다.", size=10, color=ORANGE)
        tbl(['네임스페이스', '대상', '이유', '횟수', '메시지'],
            [[e.get('namespace',''), e.get('obj',''), e.get('reason',''),
              str(e.get('count',1)), e.get('message','')[:60]]
             for e in events[:20]], col_widths=[2, 3, 2, 1.5, 8])
    _top_f = '; '.join(
        f"[{sev_ko(r.get('severity',''))}] {r.get('detail', r.get('message', ''))}"
        for r in findings[:5]
    ) if findings else '없음'
    ai_comment('발견사항 종합', f"""발견 이슈 요약:
- 긴급: {ko_counts['긴급']}건, 높음: {ko_counts['높음']}건, 중간: {ko_counts['중간']}건, 낮음: {ko_counts['낮음']}건
- 주요 이슈: {_top_f}

전체 발견사항에 대한 핵심 조치 방향을 1~2문장으로 작성하세요.""")
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════════
    # 9. 예상 개선 작업 (v1 스타일 자동 생성)
    # ═══════════════════════════════════════════════════════════════════════════
    log_fn('  [10/10] 예상 개선 작업...')
    h1('9. 예상 개선 작업')
    para('발견 이슈 기반으로 자동 생성된 개선 작업 목록입니다. '
         '우선순위와 예상 공수는 이슈 유형에 따라 자동 산정되었습니다.', size=10, color=GRAY)
    spacer()

    AREA_MAP = {
        'pod_restart': '운영/SRE', 'pvc': '인프라', 'node': '인프라',
        'security': '보안', 'network': '보안/네트워크', 'rbac': '보안',
        'hpa': '운영', 'resource': '운영', 'deployment': '운영',
        'service': '네트워크', 'ingress': '네트워크',
        'helm': '운영', 'statefulset': '운영', 'daemonset': '운영',
    }
    EFFORT_MAP = {'긴급': '1~3일', '높음': '1~5일', '중간': '3~10일', '낮음': '1~2주'}

    work_rows = []
    seen = set()
    for r in (priority['immediate'] + priority['short_term'] + (priority.get('backlog') or [])):
        rec = r.get('recommendation', '')[:60]
        if rec in seen: continue
        seen.add(rec)
        cat  = r.get('category', '').lower()
        area = next((v for k, v in AREA_MAP.items() if k in cat), '운영')
        sev  = sev_ko(r['severity'])
        work_rows.append([
            str(len(work_rows) + 1),
            rec,
            area,
            EFFORT_MAP.get(sev, '미정'),
            (sev, True, sev_rgb(r['severity']), 'CENTER', sev_bg(r['severity'])),
        ])
        if len(work_rows) >= 20: break

    if work_rows:
        tbl(['#', '작업명', '담당 영역', '예상 공수', '우선순위'],
            work_rows, col_widths=[0.8, 8, 2.5, 2, 3.2])
    else:
        callout('✓ 즉시 개선 작업이 필요한 이슈가 없습니다.', HEX_LOWD, GREEN)


    _fix_ea(doc)
    doc.save(out_path)
    return out_path


def _report_write_txt(data, findings, llm_summary, out_path):
    """python-docx 없을 때 TXT 폴백."""
    priority = _report_build_priority_summary(findings)
    lines = [
        '=' * 60,
        'Kubernetes 클러스터 보고서',
        f"생성: {data.get('collected_at', '')}",
        f"버전: {data.get('cluster_version', 'N/A')}",
        '=' * 60, '',
        f"노드: {len(data.get('nodes') or [])}, "
        f"파드: {len(data.get('pods') or [])}, "
        f"Deployment: {len(data.get('deployments') or [])}, "
        f"Job: {len(data.get('jobs') or [])}, "
        f"CronJob: {len(data.get('cronjobs') or [])}",
        '', '── 발견 사항 ──',
    ]
    for f in findings:
        lines.append(
            f"  [{(f.get('severity') or f.get('level', 'INFO')).upper()}] "
            f"{f.get('namespace', 'cluster')}/{f.get('name', '-')}: "
            f"{f.get('value', '-')} — {f.get('detail') or f.get('message', '')} "
            f"(권장: {f.get('rec', '-')})"
        )
    lines += ['', '── 개선 우선순위 ──']
    lines.append(
        f"긴급 {priority['counts'].get('critical', 0)} / "
        f"높음 {priority['counts'].get('high', 0)} / "
        f"중간 {priority['counts'].get('medium', 0)}"
    )
    for row in priority['immediate'][:10]:
        lines.append(f"  [즉시] {row['target']} — {row['recommendation']}")
    for row in priority['short_term'][:10]:
        lines.append(f"  [단기] {row['target']} — {row['recommendation']}")
    if data.get('events'):
        lines += ['', '── Warning 이벤트 ──']
        for e in (data.get('events') or [])[:30]:
            lines.append(
                f"  [{e.get('reason', '-')}] {e.get('namespace', '')}/{e.get('obj', '')}: "
                f"{e.get('message', '')}"
            )
    if llm_summary:
        lines += ['', '── AI 분석 ──', llm_summary]
    Path(out_path).write_text('\n'.join(lines), encoding='utf-8')
    return out_path




# ─────────────────────────────────────────────────────────────────────────────

def _esc(s):
    return str(s).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')


def _table(headers: list, rows: list, keys: list) -> str:
    th = ''.join(f'<th>{h}</th>' for h in headers)
    trs = ''
    for row in rows:
        tds = ''.join(f'<td>{_esc(row.get(k, ""))}</td>' for k in keys)
        trs += f'<tr>{tds}</tr>'
    return f'<table><thead><tr>{th}</tr></thead><tbody>{trs}</tbody></table>'


def _build_report_html(k8s: K8sManager, ns: str, sections: set) -> str:
    now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    ver = k8s.cluster_info.get('version', 'N/A')
    parts = []

    def section(title, content):
        parts.append(f'<section><h2>{_esc(title)}</h2>{content}</section>')

    # ── 노드 ──
    if 'nodes' in sections:
        rows = k8s.get_nodes()
        section('노드 상태', _table(
            ['이름', '상태', '역할', 'Version', 'OS', 'Age'],
            rows, ['name', 'status', 'roles', 'version', 'os', 'age']))

    # ── 파드 ──
    if 'pods' in sections:
        rows = [{k: v for k, v in r.items() if not k.startswith('_')}
                for r in k8s.get_pods(ns or None)]
        section(f'파드{"" if not ns else f" ({ns})"}', _table(
            ['이름', '네임스페이스', '상태', 'Ready', '재시작', '노드', 'Age'],
            rows, ['name', 'namespace', 'status', 'ready', 'restarts', 'node', 'age']))

    # ── 디플로이먼트 ──
    if 'deployments' in sections:
        rows = [{k: v for k, v in r.items() if not k.startswith('_')}
                for r in k8s.get_deployments(ns or None)]
        section(f'디플로이먼트', _table(
            ['이름', '네임스페이스', 'Ready', 'Up-to-date', 'Available', 'Age'],
            rows, ['name', 'namespace', 'ready', 'up-to-date', 'available', 'age']))

    # ── 서비스 ──
    if 'services' in sections:
        rows = [{k: v for k, v in r.items() if not k.startswith('_')}
                for r in k8s.get_services(ns or None)]
        section('서비스', _table(
            ['이름', '네임스페이스', '타입', 'Cluster IP', 'External IP', 'Port(s)', 'Age'],
            rows, ['name', 'namespace', 'type', 'cluster-ip', 'external-ip', 'ports', 'age']))

    # ── PVC ──
    if 'pvcs' in sections:
        rows = [{k: v for k, v in r.items() if not k.startswith('_')}
                for r in k8s.get_pvcs(ns or None)]
        section('PVC', _table(
            ['이름', '네임스페이스', '상태', 'Volume', 'Capacity', 'Access Modes', 'Age'],
            rows, ['name', 'namespace', 'status', 'volume', 'capacity', 'access-modes', 'age']))

    # ── Helm ──
    if 'helm' in sections:
        rows = [{k: v for k, v in r.items() if not k.startswith('_')}
                for r in k8s.get_helm_releases()]
        section('Helm 릴리스', _table(
            ['이름', '네임스페이스', '차트', 'App Version', '상태', 'Revision', '업데이트'],
            rows, ['name', 'namespace', 'chart', 'app_version', 'status', 'revision', 'updated']))

    # ── ArgoCD ──
    if 'argocd' in sections:
        raw = k8s.get_argocd_apps()
        if raw:
            rows = [PolarisAPI._transform_argo(item) for item in raw]
            section('ArgoCD 앱', _table(
                ['이름', '네임스페이스', 'Sync', 'Health', 'Repo', 'Revision'],
                rows, ['name', 'namespace', 'sync', 'health', 'repo', 'revision']))

    body = '\n'.join(parts) if parts else '<p>선택된 섹션이 없습니다.</p>'

    return f'''<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="UTF-8">
<title>Polaris 클러스터 보고서</title>
<style>
  body {{ font-family: "Segoe UI", "Noto Sans KR", sans-serif; margin: 0; padding: 24px;
         background: #f8fafc; color: #1e293b; font-size: 13px; }}
  h1 {{ font-size: 22px; color: #0f172a; margin-bottom: 4px; }}
  .meta {{ color: #64748b; font-size: 12px; margin-bottom: 32px; }}
  section {{ margin-bottom: 32px; }}
  h2 {{ font-size: 14px; font-weight: 700; color: #1e40af; border-bottom: 2px solid #bfdbfe;
        padding-bottom: 6px; margin-bottom: 12px; }}
  table {{ width: 100%; border-collapse: collapse; background: #fff;
           box-shadow: 0 1px 3px rgba(0,0,0,.08); border-radius: 6px; overflow: hidden; }}
  thead {{ background: #1e3a5f; color: #fff; }}
  th {{ padding: 9px 12px; text-align: left; font-size: 11px; font-weight: 600;
        text-transform: uppercase; letter-spacing: .05em; }}
  td {{ padding: 7px 12px; border-bottom: 1px solid #e2e8f0; }}
  tr:last-child td {{ border-bottom: none; }}
  tr:hover td {{ background: #f1f5f9; }}
</style>
</head>
<body>
<h1>⎈ Polaris 클러스터 보고서</h1>
<div class="meta">생성 시각: {now} &nbsp;·&nbsp; 클러스터 버전: {_esc(ver)}</div>
{body}
</body>
</html>'''


